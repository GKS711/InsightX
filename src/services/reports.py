"""
v6 Report generator — PDF (reportlab) + DOCX (python-docx). Sync.

Workflow:
  1. caller posts to POST /api/v5/stores/{id}/reports → builds report job row
  2. report_job_runner gathers latest succeeded analysis_runs (analyze + swot + reply + weekly_plan)
  3. composes PDF or DOCX, writes to outputs/reports/store_{id}_{ts}.{ext}
  4. updates Report row with file_path
"""
from __future__ import annotations

import json
import logging
import os
from datetime import datetime, timezone
from typing import Any

from sqlalchemy import select
from sqlalchemy.orm import Session

from src.models import AnalysisRun, Report, Store

logger = logging.getLogger(__name__)

REPORTS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "outputs",
    "reports",
)


def gather_report_inputs(
    session: Session, store_id: int
) -> dict[str, Any]:
    """Collect latest succeeded run per ai_function for the store."""
    out: dict[str, Any] = {}
    for fn in ["analyze", "swot", "reply", "weekly_plan", "internal_email"]:
        stmt = (
            select(AnalysisRun)
            .where(
                AnalysisRun.store_id == store_id,
                AnalysisRun.ai_function == fn,
                AnalysisRun.status == "succeeded",
            )
            .order_by(AnalysisRun.created_at.desc())
            .limit(1)
        )
        result = session.execute(stmt)
        run = result.scalar_one_or_none()
        if run is not None:
            out[fn] = run.output_json
    return out


def build_pdf(store: Store, inputs: dict[str, Any], file_path: str) -> None:
    """Build a PDF report using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.lib import colors

    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title", parent=styles["Title"], fontSize=22, leading=26, spaceAfter=18,
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontSize=14, leading=18, spaceAfter=10, spaceBefore=14,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["BodyText"], fontSize=10, leading=14,
    )

    story: list[Any] = []

    # Title
    story.append(Paragraph(f"<b>InsightX 週報：{store.name}</b>", title_style))
    story.append(
        Paragraph(
            f"產生時間：{datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
            body_style,
        )
    )
    story.append(Spacer(1, 12))

    # Section: analyze (good/bad themes)
    analyze = inputs.get("analyze") or {}
    if analyze:
        story.append(Paragraph("一、整體分析", h2_style))
        if analyze.get("total_reviews"):
            story.append(Paragraph(f"資料規模：{analyze['total_reviews']}", body_style))

        good = analyze.get("good") or []
        bad = analyze.get("bad") or []
        if good or bad:
            data = [["優點 / 占比", "改善點 / 占比"]]
            max_len = max(len(good), len(bad))
            for i in range(max_len):
                g = good[i] if i < len(good) else {}
                b = bad[i] if i < len(bad) else {}
                data.append(
                    [
                        f"{g.get('label', '-')}（{g.get('value', '-')}%）" if g else "—",
                        f"{b.get('label', '-')}（{b.get('value', '-')}%）" if b else "—",
                    ]
                )
            tbl = Table(data, colWidths=[8 * cm, 8 * cm])
            tbl.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#5e6ad2")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTSIZE", (0, 0), (-1, -1), 10),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
                        ("BOX", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ]
                )
            )
            story.append(tbl)
            story.append(Spacer(1, 12))

    # Section: SWOT
    swot = inputs.get("swot") or {}
    if swot and (swot.get("strengths") or swot.get("weaknesses")):
        story.append(Paragraph("二、SWOT 分析", h2_style))
        for label_en, label_zh in [
            ("strengths", "優勢 (S)"),
            ("weaknesses", "劣勢 (W)"),
            ("opportunities", "機會 (O)"),
            ("threats", "威脅 (T)"),
        ]:
            items = swot.get(label_en) or []
            if items:
                story.append(Paragraph(f"<b>{label_zh}</b>", body_style))
                for it in items[:5]:
                    if isinstance(it, dict):
                        text = it.get("text") or it.get("label") or str(it)
                    else:
                        text = str(it)
                    story.append(Paragraph(f"• {text}", body_style))
                story.append(Spacer(1, 6))

    # Section: weekly plan
    wp = inputs.get("weekly_plan")
    if wp:
        story.append(Paragraph("三、本週行動計畫", h2_style))
        if isinstance(wp, dict):
            text = wp.get("text") or json.dumps(wp, ensure_ascii=False, indent=2)
        else:
            text = str(wp)
        for line in text.split("\n")[:30]:
            line = line.strip()
            if line:
                story.append(Paragraph(line, body_style))

    # Section: internal email
    email = inputs.get("internal_email")
    if email:
        story.append(Paragraph("四、內部團隊週報信", h2_style))
        if isinstance(email, dict):
            text = email.get("text") or json.dumps(email, ensure_ascii=False, indent=2)
        else:
            text = str(email)
        for line in text.split("\n")[:30]:
            line = line.strip()
            if line:
                story.append(Paragraph(line, body_style))

    if len(story) < 3:
        story.append(Paragraph("（尚無足夠分析資料 — 請先觸發 analyze / swot / weekly_plan run）", body_style))

    doc.build(story)


def build_docx(store: Store, inputs: dict[str, Any], file_path: str) -> None:
    """Build a DOCX report using python-docx."""
    from docx import Document

    doc = Document()
    doc.add_heading(f"InsightX 週報：{store.name}", level=0)
    doc.add_paragraph(
        f"產生時間：{datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )

    analyze = inputs.get("analyze") or {}
    if analyze:
        doc.add_heading("一、整體分析", level=1)
        if analyze.get("total_reviews"):
            doc.add_paragraph(f"資料規模：{analyze['total_reviews']}")
        good = analyze.get("good") or []
        bad = analyze.get("bad") or []
        if good or bad:
            tbl = doc.add_table(rows=1 + max(len(good), len(bad)), cols=2)
            tbl.style = "Light List Accent 1"
            tbl.cell(0, 0).text = "優點"
            tbl.cell(0, 1).text = "改善點"
            for i in range(max(len(good), len(bad))):
                if i < len(good):
                    g = good[i]
                    tbl.cell(i + 1, 0).text = f"{g.get('label', '-')} ({g.get('value', '-')}%)"
                if i < len(bad):
                    b = bad[i]
                    tbl.cell(i + 1, 1).text = f"{b.get('label', '-')} ({b.get('value', '-')}%)"

    swot = inputs.get("swot") or {}
    if swot and (swot.get("strengths") or swot.get("weaknesses")):
        doc.add_heading("二、SWOT 分析", level=1)
        for label_en, label_zh in [
            ("strengths", "優勢 (S)"),
            ("weaknesses", "劣勢 (W)"),
            ("opportunities", "機會 (O)"),
            ("threats", "威脅 (T)"),
        ]:
            items = swot.get(label_en) or []
            if items:
                doc.add_heading(label_zh, level=2)
                for it in items[:5]:
                    text = it.get("text") if isinstance(it, dict) else str(it)
                    doc.add_paragraph(text or "—", style="List Bullet")

    wp = inputs.get("weekly_plan")
    if wp:
        doc.add_heading("三、本週行動計畫", level=1)
        text = wp.get("text") if isinstance(wp, dict) else str(wp)
        for line in (text or "").split("\n")[:30]:
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    if len(doc.paragraphs) < 3:
        doc.add_paragraph("（尚無足夠分析資料）")

    doc.save(file_path)


def generate_report(
    session: Session,
    store: Store,
    report: Report,
) -> None:
    """Synchronously build the report file and update the Report row."""
    os.makedirs(REPORTS_DIR, exist_ok=True)
    inputs = gather_report_inputs(session, store.id)

    ts = datetime.now(tz=timezone.utc).strftime("%Y%m%dT%H%M%S")
    ext = report.format
    fname = f"store{store.id}_{ts}.{ext}"
    fpath = os.path.join(REPORTS_DIR, fname)

    try:
        if ext == "pdf":
            build_pdf(store, inputs, fpath)
        elif ext == "docx":
            build_docx(store, inputs, fpath)
        else:
            raise ValueError(f"unsupported format: {ext}")

        report.status = "succeeded"
        report.file_path = fpath
        report.generated_at = datetime.now(tz=timezone.utc)
    except Exception as exc:
        logger.warning("[reports] generate failed: %s", exc)
        report.status = "failed"
        report.error_message = str(exc)[:1000]
